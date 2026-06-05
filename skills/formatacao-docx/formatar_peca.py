#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Biblioteca de formatação de peças jurídicas da DPU.
Converte texto estruturado em .docx formatado + PDF.

Versão 2.0: Suporte a notas de rodapé via Banco de Fontes Verificadas.
Marcadores [Fxxx] no texto são convertidos em notas de rodapé reais no DOCX
quando o parâmetro --banco é fornecido.
"""

import os
import sys
import re
import json
import platform

# Forçar stdout/stderr para UTF-8 no Windows (evita erros de encoding no terminal)
if sys.stdout.encoding != 'utf-8':
    sys.stdout = open(sys.stdout.fileno(), mode='w', encoding='utf-8', buffering=1)
if sys.stderr.encoding != 'utf-8':
    sys.stderr = open(sys.stderr.fileno(), mode='w', encoding='utf-8', buffering=1)

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from lxml import etree

# Caminhos padrão (env vars overrideiam)
_SKILL_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_PATH = os.environ.get(
    "FORMATAR_PECA_TEMPLATE",
    os.path.join(_SKILL_DIR, "assets", "template_dpu.docx"),
)
SAIDA_DIR = os.environ.get("FORMATAR_PECA_SAIDA_DIR", os.path.join(os.getcwd(), "saida"))

# Assinatura (parametrizável — cada defensor configura a sua via env var ou argumento)
ASSINATURA_NOME = os.environ.get("FORMATAR_PECA_NOME", "[NOME DO(A) DEFENSOR(A)]")
ASSINATURA_CARGO = os.environ.get("FORMATAR_PECA_CARGO", "Defensor(a) Público(a) Federal")
ASSINATURA_LOCAL = os.environ.get("FORMATAR_PECA_LOCAL", "[CIDADE/UF]")


def atualizar_barra_lateral(doc, texto_pedido, itens_sumario):
    """
    Atualiza o conteúdo da barra lateral na primeira página (header2).
    OBRIGATÓRIO ao usar o template.
    """
    section = doc.sections[0]
    first_header = section.first_page_header
    header_element = first_header._element

    wps_ns = 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape'
    for txbx in header_element.iter('{%s}txbx' % wps_ns):
        txbxContent = txbx.find(qn('w:txbxContent'))
        if txbxContent is None:
            continue

        # Verificar se é a textbox da barra lateral
        full_text = ''.join(t.text or '' for t in txbxContent.iter(qn('w:t')))
        if 'Pedido' not in full_text and 'Sum' not in full_text:
            continue

        # Limpar todo o conteúdo existente
        for child in list(txbxContent):
            txbxContent.remove(child)

        # Título "Pedido"
        txbxContent.append(_criar_paragrafo_barra(
            texto="Pedido", fonte="Quattrocento Sans", tamanho=28,
            negrito=True, italico=False, cor="323e4f",
            space_before=120, space_after=240, line_spacing="276"))

        # Texto do pedido
        txbxContent.append(_criar_paragrafo_barra(
            texto=texto_pedido, fonte="Palatino Linotype", tamanho=22,
            negrito=False, italico=True, cor="323e4f",
            space_before=120, space_after=0, line_spacing="276"))

        # Título "Sumário"
        txbxContent.append(_criar_paragrafo_barra(
            texto="Sumário", fonte="Quattrocento Sans", tamanho=28,
            negrito=True, italico=False, cor="323e4f",
            space_before=240, space_after=0, line_spacing="240"))

        # Itens do sumário
        for item in itens_sumario:
            txbxContent.append(_criar_paragrafo_barra(
                texto=item, fonte="Quattrocento Sans", tamanho=18,
                negrito=False, italico=False, cor="323e4f",
                space_before=60, space_after=60, line_spacing="240"))

        break


def _criar_paragrafo_barra(texto, fonte, tamanho, negrito, italico, cor,
                           space_before, space_after, line_spacing):
    """Cria elemento w:p para inserir na textbox da barra lateral."""
    p = etree.SubElement(etree.Element('dummy'), qn('w:p'))
    p.getparent().remove(p)

    pPr = etree.SubElement(p, qn('w:pPr'))
    spacing = etree.SubElement(pPr, qn('w:spacing'))
    spacing.set(qn('w:after'), str(space_after))
    spacing.set(qn('w:before'), str(space_before))
    spacing.set(qn('w:line'), str(line_spacing))
    ind = etree.SubElement(pPr, qn('w:ind'))
    ind.set(qn('w:left'), '0')
    ind.set(qn('w:right'), '0')
    ind.set(qn('w:firstLine'), '0')
    jc = etree.SubElement(pPr, qn('w:jc'))
    jc.set(qn('w:val'), 'both')
    td = etree.SubElement(pPr, qn('w:textDirection'))
    td.set(qn('w:val'), 'btLr')

    r = etree.SubElement(p, qn('w:r'))
    rPr = etree.SubElement(r, qn('w:rPr'))
    rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:ascii'), fonte)
    rFonts.set(qn('w:cs'), fonte)
    rFonts.set(qn('w:eastAsia'), fonte)
    rFonts.set(qn('w:hAnsi'), fonte)
    b = etree.SubElement(rPr, qn('w:b'))
    b.set(qn('w:val'), '1' if negrito else '0')
    i = etree.SubElement(rPr, qn('w:i'))
    i.set(qn('w:val'), '1' if italico else '0')
    smallCaps = etree.SubElement(rPr, qn('w:smallCaps'))
    smallCaps.set(qn('w:val'), '0')
    strike = etree.SubElement(rPr, qn('w:strike'))
    strike.set(qn('w:val'), '0')
    color = etree.SubElement(rPr, qn('w:color'))
    color.set(qn('w:val'), cor)
    sz = etree.SubElement(rPr, qn('w:sz'))
    sz.set(qn('w:val'), str(tamanho))
    vertAlign = etree.SubElement(rPr, qn('w:vertAlign'))
    vertAlign.set(qn('w:val'), 'baseline')
    t = etree.SubElement(r, qn('w:t'))
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = texto

    return p


def atualizar_rodape_primeira_pagina(doc, texto_rodape):
    """Substitui o texto do rodapé da primeira página (footer2)."""
    section = doc.sections[0]
    first_footer = section.first_page_footer
    footer_element = first_footer._element

    for t_elem in footer_element.iter(qn('w:t')):
        if t_elem.text and 'PAJ' in t_elem.text:
            t_elem.text = texto_rodape
            return

    # Fallback: substituir primeiro texto encontrado
    for t_elem in footer_element.iter(qn('w:t')):
        if t_elem.text and t_elem.text.strip():
            t_elem.text = texto_rodape
            return


def carregar_banco_fontes(banco_path):
    """Carrega o Banco de Fontes Verificadas (JSON)."""
    if not banco_path or not os.path.exists(banco_path):
        return None
    try:
        with open(banco_path, 'r', encoding='utf-8') as f:
            banco = json.load(f)
        fontes_map = {}
        for fonte in banco.get('fontes', []):
            fontes_map[fonte['id']] = fonte
        return fontes_map
    except Exception as e:
        print(f"[AVISO] Erro ao carregar Banco de Fontes: {e}")
        return None


def _formatar_texto_nota_rodape(fonte):
    """Formata o texto da nota de rodapé a partir de uma entrada do Banco."""
    citacao = fonte.get('citacao', '')
    verif = fonte.get('verificacao', {})
    metodo = verif.get('metodo', '')

    texto = citacao
    if metodo == 'url':
        texto += f". Disponível em: {verif.get('url', '')}"
    elif metodo == 'documento':
        doc = verif.get('documento', '')
        pag = verif.get('pagina', '')
        texto += f". Fonte: {doc}, p. {pag}"

    return texto


def _criar_nota_rodape_xml(doc, texto_nota, nota_id):
    """
    Cria uma nota de rodapé real no DOCX usando manipulação XML direta.
    Retorna o elemento w:r com a referência à nota.
    """
    # Acessar ou criar a parte de footnotes
    footnotes_part = None
    for rel in doc.part.rels.values():
        if "footnotes" in rel.reltype:
            footnotes_part = rel.target_part
            break

    if footnotes_part is None:
        # Se não existe parte de footnotes, criar nota simples inline
        return None

    # Acessar o XML — Part genérico usa .blob, XmlPart usa ._element
    if hasattr(footnotes_part, '_element'):
        footnotes_xml = footnotes_part._element
    else:
        # Part genérico: parsear o blob XML e reatribuir
        if not hasattr(footnotes_part, '_footnotes_xml_cache'):
            footnotes_part._footnotes_xml_cache = etree.fromstring(footnotes_part.blob)
        footnotes_xml = footnotes_part._footnotes_xml_cache
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    footnote = etree.SubElement(footnotes_xml, qn('w:footnote'))
    footnote.set(qn('w:id'), str(nota_id))

    # Parágrafo da nota
    p = etree.SubElement(footnote, qn('w:p'))
    pPr = etree.SubElement(p, qn('w:pPr'))
    pStyle = etree.SubElement(pPr, qn('w:pStyle'))
    pStyle.set(qn('w:val'), 'FootnoteText')

    # Run com referência da nota
    r_ref = etree.SubElement(p, qn('w:r'))
    rPr_ref = etree.SubElement(r_ref, qn('w:rPr'))
    rStyle = etree.SubElement(rPr_ref, qn('w:rStyle'))
    rStyle.set(qn('w:val'), 'FootnoteReference')
    fn_ref = etree.SubElement(r_ref, qn('w:footnoteRef'))

    # Run com texto da nota
    r_text = etree.SubElement(p, qn('w:r'))
    rPr_text = etree.SubElement(r_text, qn('w:rPr'))
    rFonts = etree.SubElement(rPr_text, qn('w:rFonts'))
    rFonts.set(qn('w:ascii'), 'Quattrocento Sans')
    rFonts.set(qn('w:hAnsi'), 'Quattrocento Sans')
    sz = etree.SubElement(rPr_text, qn('w:sz'))
    sz.set(qn('w:val'), '16')  # 8pt
    t = etree.SubElement(r_text, qn('w:t'))
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = f" {texto_nota}"

    return nota_id


def _inserir_referencia_nota(paragrafo_element, nota_id, posicao_run=None):
    """
    Insere referência de nota de rodapé (número sobrescrito) no parágrafo.
    Retorna o run XML criado.
    """
    r = etree.SubElement(paragrafo_element, qn('w:r'))
    rPr = etree.SubElement(r, qn('w:rPr'))
    rStyle = etree.SubElement(rPr, qn('w:rStyle'))
    rStyle.set(qn('w:val'), 'FootnoteReference')
    fn_ref = etree.SubElement(r, qn('w:footnoteReference'))
    fn_ref.set(qn('w:id'), str(nota_id))
    return r


def _numero_sobrescrito(n):
    """Converte número inteiro em string Unicode sobrescrita."""
    superscripts = {'0': '\u2070', '1': '\u00b9', '2': '\u00b2', '3': '\u00b3',
                    '4': '\u2074', '5': '\u2075', '6': '\u2076', '7': '\u2077',
                    '8': '\u2078', '9': '\u2079'}
    return ''.join(superscripts[d] for d in str(n))


def processar_notas_rodape(doc, texto, banco_fontes):
    """
    Processa marcadores [Fxxx] no texto e prepara mapeamento para notas de rodapé.
    Substitui marcadores por números sobrescritos Unicode e gera seção de notas ao final.
    Retorna: (texto_com_notas, lista_notas)
    """
    if not banco_fontes:
        return texto, []

    pattern = r'\[F(\d{3})\]'
    notas = []
    contador_nota = 1

    # Encontrar todos os marcadores e suas fontes
    for match in re.finditer(pattern, texto):
        id_fonte = f"F{match.group(1)}"
        if id_fonte in banco_fontes:
            fonte = banco_fontes[id_fonte]
            texto_nota = _formatar_texto_nota_rodape(fonte)

            notas.append({
                'id_fonte': id_fonte,
                'nota_id': contador_nota,
                'texto_nota': texto_nota,
                'marcador': match.group(0),
            })
            contador_nota += 1

    # Substituir marcadores por números sobrescritos Unicode
    nota_idx = [0]  # usar lista para closure mutável
    def _substituir(match):
        id_fonte = f"F{match.group(1)}"
        if id_fonte in banco_fontes:
            # Encontrar a nota correspondente
            for nota in notas:
                if nota['id_fonte'] == id_fonte:
                    return _numero_sobrescrito(nota['nota_id'])
        return ''
    texto_com_notas = re.sub(pattern, _substituir, texto)

    # Notas serão adicionadas APÓS a assinatura pelo criar_documento
    return texto_com_notas, notas


def criar_documento(
    conteudo_txt,
    enderecamento,
    titulo_peca,
    subtitulo=None,
    numero_paj=None,
    numero_processo=None,
    texto_pedido=None,
    itens_sumario=None,
    template_path=None,
    banco_path=None,
    assinatura_nome=None,
    assinatura_cargo=None,
    assinatura_local=None
):
    """
    Cria documento .docx formatado a partir de texto estruturado.

    Args:
        conteudo_txt: str - Texto completo da peça com marcadores:
            ## Título Principal (Heading 1)
            ### Subtítulo (Heading 3)
            > Citação longa (recuo 4cm, fonte 10pt)
            [Fxxx] Marcador de fonte (convertido em nota de rodapé se banco fornecido)
            Texto normal (parágrafo com recuo de primeira linha)

        enderecamento: str - "EXCELENTÍSSIMO SENHOR..."
        titulo_peca: str - "AGRAVO INTERNO"
        subtitulo: str (opcional) - Subtítulo abaixo do título
        numero_paj: str (opcional) - "PAJ XXXX/XXX-XXXXX"
        numero_processo: str (opcional) - Usado no rodapé se não houver PAJ
        texto_pedido: str (opcional) - Resumo para barra lateral
        itens_sumario: list (opcional) - Itens do sumário para barra lateral
        template_path: str (opcional) - Caminho customizado do template
        banco_path: str (opcional) - Caminho do Banco de Fontes Verificadas (JSON).
            Se fornecido, marcadores [Fxxx] são convertidos em notas de rodapé.

    Returns:
        Document - Documento python-docx pronto para salvar
    """
    if template_path is None:
        template_path = TEMPLATE_PATH

    print("Carregando template...")
    doc = Document(template_path)

    # Carregar Banco de Fontes e processar notas de rodapé
    banco_fontes = None
    notas_rodape = []
    if banco_path:
        banco_fontes = carregar_banco_fontes(banco_path)
        if banco_fontes:
            print(f"[OK] Banco de Fontes carregado: {len(banco_fontes)} fontes")
            conteudo_txt, notas_rodape = processar_notas_rodape(doc, conteudo_txt, banco_fontes)
            if notas_rodape:
                print(f"[OK] {len(notas_rodape)} notas de verificação inseridas no texto")

    # Limpar conteúdo do corpo, preservando seção
    body = doc.element.body
    for child in list(body):
        if child.tag != qn('w:sectPr'):
            body.remove(child)

    # OBRIGATÓRIO: Atualizar barra lateral — SEMPRE, para não deixar texto do modelo
    print("Atualizando barra lateral...")
    texto_pedido_barra = texto_pedido or _extrair_texto_pedido(conteudo_txt)
    itens_sumario_barra = itens_sumario or _extrair_sumario(conteudo_txt) or ['(sem sumário)']
    atualizar_barra_lateral(doc, texto_pedido_barra, itens_sumario_barra)

    # OBRIGATÓRIO: Atualizar rodapé
    print("Atualizando rodapé...")
    texto_rodape = numero_paj if numero_paj else numero_processo
    if texto_rodape:
        atualizar_rodape_primeira_pagina(doc, texto_rodape)

    # Processar conteúdo
    print("Processando conteúdo...")
    paragrafos = _processar_texto(conteudo_txt)

    # Endereçamento — suporta múltiplas linhas (cada linha = parágrafo em negrito)
    linhas_end = enderecamento.split('\n') if enderecamento else ['']
    for i, linha in enumerate(linhas_end):
        p_end = doc.add_paragraph()
        p_end.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_end.paragraph_format.space_before = Pt(0)
        p_end.paragraph_format.space_after = Pt(30) if i == len(linhas_end) - 1 else Pt(2)
        p_end.paragraph_format.first_line_indent = Pt(0)
        run = p_end.add_run(linha)
        run.bold = True
        run.font.name = 'Quattrocento Sans'
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)

    # Título da peça
    p_tit = doc.add_paragraph()
    p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_tit.paragraph_format.space_before = Pt(12)
    p_tit.paragraph_format.space_after = Pt(24) if not subtitulo else Pt(6)
    p_tit.paragraph_format.first_line_indent = Pt(0)
    run = p_tit.add_run(titulo_peca)
    run.bold = True
    run.font.name = 'Quattrocento Sans'
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 0, 0)

    # Subtítulo (se houver)
    if subtitulo:
        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sub.paragraph_format.space_before = Pt(0)
        p_sub.paragraph_format.space_after = Pt(36)
        p_sub.paragraph_format.first_line_indent = Pt(0)
        run = p_sub.add_run(subtitulo)
        run.font.name = 'Quattrocento Sans'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    # Adicionar parágrafos
    print("Inserindo conteúdo...")
    for item in paragrafos:
        _adicionar_paragrafo(doc, item)

    # Assinatura
    print("Adicionando assinatura...")
    _adicionar_assinatura(doc, assinatura_nome, assinatura_cargo, assinatura_local)

    # Seção de citações (após assinatura, em nova página)
    if notas_rodape:
        print("Adicionando seção de citações...")
        _adicionar_secao_citacoes(doc, notas_rodape)

    return doc


def _processar_texto(texto):
    """
    Processa texto com marcadores e retorna lista de parágrafos estruturados.

    Marcadores suportados:
        ## Título Principal → heading1
        ### Subtítulo → heading3
        > Citação longa → citacao_longa
        **texto** → negrito inline (dentro de parágrafos normais e citações)
        Texto normal → normal

    IMPORTANTE:
        - O texto de entrada DEVE usar estes marcadores específicos
        - **negrito** funciona inline em parágrafos normais e citações longas
        - *itálico* NÃO é suportado
        - Salve o arquivo .txt com encoding='utf-8' para evitar corrupção de acentos
        - Use linhas em branco para separar parágrafos

    Exemplo de texto correto:
        ## I – INTRODUÇÃO

        Primeiro parágrafo.

        Segundo parágrafo.

        ### Subtópico

        Texto do subtópico.

        > Citação longa com recuo.

        Texto normal.
    """
    linhas = texto.split('\n')
    paragrafos = []
    buffer = []
    tipo_atual = 'normal'

    for linha in linhas:
        linha_stripped = linha.strip()

        if linha_stripped.startswith('## '):
            # Heading 1
            if buffer:
                paragrafos.append({'texto': ' '.join(buffer), 'tipo': tipo_atual})
                buffer = []
            texto = linha_stripped[3:].strip()
            paragrafos.append({'texto': texto, 'tipo': 'heading1'})
            tipo_atual = 'normal'

        elif linha_stripped.startswith('### '):
            # Heading 3
            if buffer:
                paragrafos.append({'texto': ' '.join(buffer), 'tipo': tipo_atual})
                buffer = []
            texto = linha_stripped[4:].strip()
            paragrafos.append({'texto': texto, 'tipo': 'heading3'})
            tipo_atual = 'normal'

        elif linha_stripped.startswith('>'):
            # Citação longa
            if buffer and tipo_atual != 'citacao_longa':
                paragrafos.append({'texto': ' '.join(buffer), 'tipo': tipo_atual})
                buffer = []
            texto = linha_stripped[1:].strip()
            if texto:
                buffer.append(texto)
            tipo_atual = 'citacao_longa'

        elif linha_stripped == '':
            # Linha vazia - finalizar parágrafo
            if buffer:
                paragrafos.append({'texto': ' '.join(buffer), 'tipo': tipo_atual})
                buffer = []
            tipo_atual = 'normal'

        else:
            # Linha de texto normal
            buffer.append(linha.rstrip())

    # Finalizar último parágrafo
    if buffer:
        paragrafos.append({'texto': ' '.join(buffer), 'tipo': tipo_atual})

    return paragrafos


def _adicionar_runs_com_negrito(p, texto, font_name, font_size, font_color, bold_base=False):
    """
    Adiciona runs ao parágrafo interpretando marcadores **texto** como negrito inline.
    Segmentos fora de ** herdam bold_base; segmentos dentro de ** ficam em negrito.
    """
    partes = re.split(r'(\*\*.*?\*\*)', texto)
    for parte in partes:
        if parte.startswith('**') and parte.endswith('**') and len(parte) > 4:
            conteudo = parte[2:-2]
            run = p.add_run(conteudo)
            run.bold = True
        else:
            run = p.add_run(parte)
            run.bold = bold_base
        run.font.name = font_name
        run.font.size = font_size
        run.font.color.rgb = font_color


def _adicionar_paragrafo(doc, item):
    """Adiciona parágrafo ao documento conforme o tipo."""
    tipo = item['tipo']
    texto = item['texto']

    if not texto.strip():
        return

    if tipo == 'heading1':
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.first_line_indent = Pt(0)
        # Borda inferior dourada
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            '  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="BF8F00"/>'
            '</w:pBdr>'
        )
        pPr.append(pBdr)
        run = p.add_run(texto)
        run.bold = True
        run.font.name = 'Quattrocento Sans'
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x32, 0x3E, 0x4F)

    elif tipo == 'heading3':
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.first_line_indent = Pt(0)
        run = p.add_run(texto)
        run.bold = True
        run.font.name = 'Quattrocento Sans'
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x32, 0x3E, 0x4F)

    elif tipo == 'citacao_longa':
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(4)
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(6)
        _adicionar_runs_com_negrito(p, texto, 'Quattrocento Sans', Pt(10), RGBColor(0x59, 0x59, 0x59))

    else:  # normal
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.first_line_indent = Cm(2.5)
        _adicionar_runs_com_negrito(p, texto, 'Quattrocento Sans', Pt(12), RGBColor(0x59, 0x59, 0x59))


def _adicionar_assinatura(doc, nome=None, cargo=None, local=None):
    """Adiciona assinatura ao documento.

    nome/cargo/local: se None, usam as constantes do módulo (configuráveis via
    env vars FORMATAR_PECA_NOME / _CARGO / _LOCAL). Default = placeholders.
    """
    from datetime import datetime
    nome = nome or ASSINATURA_NOME
    cargo = cargo or ASSINATURA_CARGO
    local = local or ASSINATURA_LOCAL

    # Local e data
    p_local = doc.add_paragraph()
    p_local.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_local.paragraph_format.space_before = Pt(36)
    p_local.paragraph_format.space_after = Pt(12)
    p_local.paragraph_format.first_line_indent = Pt(0)
    data_hoje = datetime.now().strftime("%d de %B de %Y")
    # Traduzir mês para português
    meses = {
        'January': 'janeiro', 'February': 'fevereiro', 'March': 'março',
        'April': 'abril', 'May': 'maio', 'June': 'junho',
        'July': 'julho', 'August': 'agosto', 'September': 'setembro',
        'October': 'outubro', 'November': 'novembro', 'December': 'dezembro'
    }
    for en, pt in meses.items():
        data_hoje = data_hoje.replace(en, pt)
    run = p_local.add_run(f"{local}, {data_hoje}.")
    run.font.name = 'Quattrocento Sans'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    # Nome
    p_assin = doc.add_paragraph()
    p_assin.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_assin.paragraph_format.space_before = Pt(48)
    p_assin.paragraph_format.space_after = Pt(0)
    p_assin.paragraph_format.first_line_indent = Pt(0)
    run = p_assin.add_run(nome.upper())
    run.bold = True
    run.font.name = 'Quattrocento Sans'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)

    # Cargo
    p_cargo = doc.add_paragraph()
    p_cargo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cargo.paragraph_format.space_before = Pt(0)
    p_cargo.paragraph_format.space_after = Pt(0)
    p_cargo.paragraph_format.first_line_indent = Pt(0)
    run = p_cargo.add_run(cargo)
    run.font.name = 'Quattrocento Sans'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)


def _adicionar_secao_citacoes(doc, notas):
    """Adiciona seção CITAÇÕES após a assinatura, em nova página."""
    from docx.enum.text import WD_BREAK

    # Quebra de página
    p_break = doc.add_paragraph()
    p_break.paragraph_format.space_before = Pt(0)
    p_break.paragraph_format.space_after = Pt(0)
    run_break = p_break.add_run()
    run_break.add_break(WD_BREAK.PAGE)

    # Título "CITAÇÕES"
    p_titulo = doc.add_paragraph()
    p_titulo.paragraph_format.space_before = Pt(18)
    p_titulo.paragraph_format.space_after = Pt(12)
    p_titulo.paragraph_format.line_spacing = 1.0
    p_titulo.paragraph_format.first_line_indent = Pt(0)
    # Borda inferior dourada (mesmo estilo dos headings)
    pPr = p_titulo._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="BF8F00"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)
    run_titulo = p_titulo.add_run("CITAÇÕES")
    run_titulo.bold = True
    run_titulo.font.name = 'Quattrocento Sans'
    run_titulo.font.size = Pt(14)
    run_titulo.font.color.rgb = RGBColor(0x32, 0x3E, 0x4F)

    # Cada nota como parágrafo
    for nota in notas:
        num = _numero_sobrescrito(nota['nota_id'])
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.first_line_indent = Pt(0)

        # Número sobrescrito
        run_num = p.add_run(f"{num} ")
        run_num.font.name = 'Quattrocento Sans'
        run_num.font.size = Pt(9)
        run_num.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

        # Texto da nota
        run_texto = p.add_run(nota['texto_nota'])
        run_texto.font.name = 'Quattrocento Sans'
        run_texto.font.size = Pt(9)
        run_texto.font.color.rgb = RGBColor(0x59, 0x59, 0x59)


def converter_para_pdf(caminho_docx):
    """
    Converte o .docx para PDF.
    Windows: usa docx2pdf (requer MS Word)
    Linux/Mac: usa LibreOffice
    """
    sistema = platform.system()
    caminho_pdf = caminho_docx.replace('.docx', '.pdf')

    if sistema == 'Windows':
        try:
            from docx2pdf import convert
            print("Convertendo para PDF (Windows)...")
            convert(caminho_docx, caminho_pdf)
            return caminho_pdf
        except ImportError:
            print("[AVISO] docx2pdf não instalado. Execute: pip install docx2pdf")
            print("[AVISO] Pulando conversão para PDF.")
            return None
    else:
        # Linux ou Mac - usar LibreOffice
        import subprocess
        print("Convertendo para PDF (LibreOffice)...")
        pasta_saida = os.path.dirname(caminho_docx)
        cmd = [
            'libreoffice',
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', pasta_saida,
            caminho_docx
        ]
        try:
            subprocess.run(cmd, check=True)
            return caminho_pdf
        except Exception as e:
            print(f"[AVISO] Erro ao converter para PDF: {e}")
            print("[AVISO] Certifique-se de que o LibreOffice está instalado.")
            return None


def salvar_documento(doc, nome_arquivo, pasta_saida=None):
    """
    Salva o documento em .docx e converte para PDF.

    Args:
        doc: Document do python-docx
        nome_arquivo: str - Nome do arquivo (ex: "agravo_interno_123456.docx")
        pasta_saida: str (opcional) - Caminho da pasta de saída

    Returns:
        tuple: (caminho_docx, caminho_pdf)
    """
    if pasta_saida is None:
        pasta_saida = SAIDA_DIR

    os.makedirs(pasta_saida, exist_ok=True)

    # Salvar .docx
    caminho_docx = os.path.join(pasta_saida, nome_arquivo)
    doc.save(caminho_docx)
    print(f"[OK] Documento DOCX salvo em: {caminho_docx}")

    # Converter para PDF
    caminho_pdf = converter_para_pdf(caminho_docx)
    if caminho_pdf:
        print(f"[OK] Documento PDF salvo em: {caminho_pdf}")

    return caminho_docx, caminho_pdf


# Mapeamento de tipo de peça para endereçamento e título
TIPOS_PECA = {
    'agravo': {
        'enderecamento_tnu': 'EXCELENTÍSSIMA TURMA NACIONAL DE UNIFORMIZAÇÃO DOS JUIZADOS ESPECIAIS FEDERAIS',
        'enderecamento_stj': 'EXCELENTÍSSIMO SENHOR MINISTRO RELATOR DO SUPERIOR TRIBUNAL DE JUSTIÇA',
        'titulo': 'AGRAVO INTERNO',
    },
    'embargos': {
        'enderecamento_tnu': 'EXCELENTÍSSIMA TURMA NACIONAL DE UNIFORMIZAÇÃO DOS JUIZADOS ESPECIAIS FEDERAIS',
        'enderecamento_stj': 'EXCELENTÍSSIMO SENHOR MINISTRO RELATOR DO SUPERIOR TRIBUNAL DE JUSTIÇA',
        'titulo': 'EMBARGOS DE DECLARAÇÃO',
    },
    'memoriais': {
        'enderecamento_tnu': 'EXCELENTÍSSIMA TURMA NACIONAL DE UNIFORMIZAÇÃO DOS JUIZADOS ESPECIAIS FEDERAIS',
        'enderecamento_stj': 'EXCELENTÍSSIMO SENHOR MINISTRO RELATOR DO SUPERIOR TRIBUNAL DE JUSTIÇA',
        'titulo': 'MEMORIAIS',
    },
    'despacho': {
        'enderecamento_tnu': '',
        'enderecamento_stj': '',
        'titulo': 'DESPACHO DE ARQUIVAMENTO',
    },
}


def _extrair_sumario(texto):
    """Extrai itens de sumário a partir dos headings ## do texto."""
    import re
    itens = []
    for match in re.finditer(r'^## (.+)$', texto, re.MULTILINE):
        itens.append(match.group(1).strip())
    return itens


def _extrair_texto_pedido(texto):
    """
    Extrai automaticamente o texto do pedido a partir do conteúdo da peça.
    Tenta: (1) parágrafo após SÍNTESE ARGUMENTATIVA, (2) último parágrafo antes de CONCLUSÃO,
    (3) primeiro parágrafo do texto. Retorna string de até 400 caracteres.
    """
    import re

    # Tentar extrair da seção SÍNTESE ARGUMENTATIVA
    m = re.search(r'SÍNTESE ARGUMENTATIVA\s*\n+(.+?)(?:\n\n|\Z)', texto, re.DOTALL)
    if m:
        trecho = m.group(1).strip()
        # Pegar apenas o primeiro parágrafo
        primeiro = trecho.split('\n\n')[0].strip()
        if len(primeiro) > 40:
            return primeiro[:400]

    # Tentar extrair do bloco CONCLUSÃO E PEDIDOS
    m = re.search(r'##[^#].*?(?:CONCLUS|PEDIDO)[^\n]*\n+(.+?)(?:\n\n|\Z)', texto, re.DOTALL | re.IGNORECASE)
    if m:
        trecho = m.group(1).strip()
        primeiro = trecho.split('\n\n')[0].strip()
        if len(primeiro) > 40:
            return primeiro[:400]

    # Fallback: primeiro parágrafo de texto normal (não título)
    for linha in texto.split('\n'):
        linha = linha.strip()
        if linha and not linha.startswith('#') and not linha.startswith('>') and len(linha) > 60:
            return linha[:400]

    return 'Recurso interposto pela Defensoria Pública da União.'


def _inferir_tribunal(processo):
    """Infere TNU ou STJ a partir do número do processo (heurística simples)."""
    if not processo:
        return 'tnu'
    # Processos STJ costumam ter segmento .3.00. ou padrão diferente
    # Por padrão, assume TNU (JEFs)
    return 'tnu'


if __name__ == '__main__':
    import argparse

    parser = argparse.ArgumentParser(
        description='Formata peça jurídica da DPU em DOCX/PDF.',
        epilog='Exemplo: python formatar_peca.py --entrada saida/agravo_VALIDADO.txt --tipo-peca agravo --numero-processo 5011388-22.2022.4.04.7200 --paj "PAJ 2026/002-10369/DPU-TNU"'
    )
    parser.add_argument('--entrada', type=str, help='Caminho do arquivo .txt validado')
    parser.add_argument('--tipo-peca', type=str, choices=list(TIPOS_PECA.keys()),
                        help='Tipo da peça: agravo, embargos, memoriais, despacho')
    parser.add_argument('--numero-processo', type=str, default=None,
                        help='Número do processo (ex: 5011388-22.2022.4.04.7200)')
    parser.add_argument('--paj', type=str, default=None,
                        help='Número do PAJ (ex: PAJ 2026/002-10369/DPU-TNU)')
    parser.add_argument('--tribunal', type=str, choices=['tnu', 'stj'], default=None,
                        help='Tribunal (tnu ou stj). Se omitido, infere do processo.')
    parser.add_argument('--banco', type=str, default=None,
                        help='Caminho do Banco de Fontes Verificadas (JSON) para gerar notas de rodapé')
    parser.add_argument('--texto-pedido', type=str, default=None,
                        help='Resumo do pedido para barra lateral (opcional)')
    parser.add_argument('--titulo', type=str, default=None,
                        help='Título da peça (substitui o título padrão do tipo). Ex: "PETIÇÃO"')
    parser.add_argument('--enderecamento', type=str, default=None,
                        help='Endereçamento completo (substitui o padrão do tipo). Suporta \\n para quebras de linha.')
    parser.add_argument('--saida', type=str, default=None,
                        help='Nome do arquivo de saída (sem extensão). Se omitido, deriva do arquivo de entrada.')
    parser.add_argument('--demo', action='store_true',
                        help='Executa formatação de demonstração com texto de exemplo')

    args = parser.parse_args()

    # Modo demo (mantém comportamento original para testes)
    if args.demo:
        conteudo = """
## I – Introdução

Este é um parágrafo normal com recuo de primeira linha. O texto segue formatado conforme o modelo institucional da DPU.

### Subtópico importante

Outro parágrafo normal. Quando citamos jurisprudência, podemos usar citações longas:

> Este é um trecho de citação longa que será formatado com recuo de 4cm à esquerda e fonte menor. Ideal para acórdãos e trechos extensos de doutrina.

## II – Conclusão

Texto final da peça antes da assinatura.
        """
        doc = criar_documento(
            conteudo_txt=conteudo,
            enderecamento="EXCELENTÍSSIMO SENHOR MINISTRO RELATOR DO SUPERIOR TRIBUNAL DE JUSTIÇA",
            titulo_peca="AGRAVO INTERNO",
            subtitulo="REsp nº 1.234.567/SP",
            numero_paj="PAJ 2024/040-12345/BVBA",
            texto_pedido="Reforma da decisão monocrática que negou seguimento ao recurso especial",
            itens_sumario=["I – Introdução", "II – Conclusão"]
        )
        salvar_documento(doc, "teste_formatacao.docx")
        print("\n[OK] Teste de demonstração concluído!")
        sys.exit(0)

    # Modo CLI normal — requer --entrada e --tipo-peca
    if not args.entrada or not args.tipo_peca:
        parser.error('Os argumentos --entrada e --tipo-peca são obrigatórios (ou use --demo para teste).')

    # Ler arquivo de entrada
    if not os.path.exists(args.entrada):
        print(f"[ERRO] Arquivo não encontrado: {args.entrada}")
        sys.exit(1)

    with open(args.entrada, 'r', encoding='utf-8') as f:
        conteudo = f.read()

    if not conteudo.strip():
        print(f"[ERRO] Arquivo vazio: {args.entrada}")
        sys.exit(1)

    # Configurar tipo de peça
    tipo = TIPOS_PECA[args.tipo_peca]
    tribunal = args.tribunal or _inferir_tribunal(args.numero_processo)
    enderecamento_key = f'enderecamento_{tribunal}'
    enderecamento = tipo.get(enderecamento_key, '')
    titulo = args.titulo if args.titulo else tipo['titulo']

    # Endereçamento: --enderecamento tem prioridade; senão usa o padrão do tipo
    if args.enderecamento:
        enderecamento = args.enderecamento.replace('\\n', '\n')
    elif args.tipo_peca == 'despacho' and args.numero_processo:
        enderecamento = f'Processo n. {args.numero_processo}'

    # Subtítulo
    subtitulo = args.paj if args.paj else None

    # Sumário automático
    itens_sumario = _extrair_sumario(conteudo)

    # Texto do pedido para barra lateral
    texto_pedido = args.texto_pedido

    # Nome do arquivo de saída
    if args.saida:
        nome_saida = args.saida if args.saida.endswith('.docx') else f"{args.saida}.docx"
    else:
        base = os.path.splitext(os.path.basename(args.entrada))[0]
        # Remover sufixo _VALIDADO se presente
        base = base.replace('_VALIDADO', '')
        nome_saida = f"{base}.docx"

    print(f"[INFO] Entrada: {args.entrada}")
    print(f"[INFO] Tipo: {args.tipo_peca} | Tribunal: {tribunal}")
    print(f"[INFO] Saída: {nome_saida}")

    # Criar documento
    doc = criar_documento(
        conteudo_txt=conteudo,
        enderecamento=enderecamento,
        titulo_peca=titulo,
        subtitulo=subtitulo,
        numero_paj=args.paj,
        numero_processo=args.numero_processo,
        texto_pedido=texto_pedido,
        itens_sumario=itens_sumario if itens_sumario else None,
        banco_path=args.banco
    )

    # Salvar
    salvar_documento(doc, nome_saida)
    print("\n[OK] Formatação concluída!")
